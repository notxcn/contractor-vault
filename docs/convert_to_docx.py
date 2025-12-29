"""
Convert Markdown to DOCX for Contractor Vault Acquisition Strategy
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def add_heading(doc, text, level):
    """Add a heading with appropriate styling."""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    
    # Headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(row_cells):
                row_cells[col_idx].text = str(cell_data)
    
    return table

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Password Manager Acquisition Landscape', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Strategic Research for Contractor Vault')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Last Updated: December 28, 2025')
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This document analyzes the password management acquisition landscape to identify what makes '
        'companies attractive acquisition targets, with specific recommendations for positioning '
        'Contractor Vault for acquisition by major players like 1Password, Bitwarden, Dashlane, or Keeper Security.'
    )
    doc.add_paragraph(
        'The password management market is rapidly consolidating around four key technology pillars: '
        'passwordless authentication, device trust/zero trust, secrets management, and SaaS access management. '
        "Contractor Vault's unique position in session sharing and contractor access management aligns closely with these trends."
    )
    
    # Part 1: Acquisition Landscape
    doc.add_heading('Part 1: Acquisition Landscape Analysis', level=1)
    
    # 1Password Acquisitions
    doc.add_heading('1Password Acquisitions (Most Active Acquirer)', level=2)
    doc.add_paragraph('1Password has made 4 strategic acquisitions since 2021:')
    
    add_table(doc, 
        ['Company', 'Date', 'Price', 'Technology', 'Strategic Rationale'],
        [
            ['Trelica', 'Jan 2025', 'Undisclosed', 'SaaS access management', 'Extend XAM platform'],
            ['Kolide', 'Feb 2024', 'Undisclosed', 'Device health & contextual access', 'Zero Trust security'],
            ['Passage Identity', 'Nov 2022', '~$24M (equity)', 'Passwordless/passkey auth', 'Enterprise passwordless'],
            ['SecretHub', 'Apr 2021', 'Undisclosed', 'Infrastructure secrets', 'DevOps credentials'],
        ]
    )
    doc.add_paragraph()
    
    # Key Insights
    doc.add_heading("Key Insights from 1Password's Strategy:", level=3)
    doc.add_paragraph('• Extended Access Management (XAM) is their flagship vision')
    doc.add_paragraph('• Focus on "user-first security" and making security accessible')
    doc.add_paragraph('• Acquiring developer-friendly tools with strong API offerings')
    doc.add_paragraph('• Building toward Zero Trust architecture')
    
    # Other Acquisitions
    doc.add_heading('Other Major Acquisitions', level=2)
    add_table(doc,
        ['Acquirer', 'Target', 'Date', 'Price', 'Technology'],
        [
            ['LogMeIn', 'LastPass', '2015', '$110M + $15M', 'Consumer password vault'],
            ['Bitwarden', 'Passwordless.dev', 'Jan 2023', 'Undisclosed', 'FIDO2/WebAuthn APIs'],
            ['Keeper Security', 'Glyptodon', 'Feb 2022', 'Undisclosed', 'Zero-trust remote access'],
            ['Okta', 'Uno', 'Oct 2023', 'Undisclosed', 'Consumer password manager'],
            ['Okta', 'Spera Security', 'Feb 2024', '~$80M', 'Identity threat detection'],
            ['Okta', 'Axiom Security', 'Aug 2025', '$100M', 'Privileged Access Management'],
        ]
    )
    doc.add_paragraph()
    
    # Part 2: What Makes Companies Attractive
    doc.add_heading('Part 2: What Makes Companies Attractive Acquisition Targets', level=1)
    
    doc.add_heading('Tier 1: Highest Value Technologies', level=2)
    add_table(doc,
        ['Technology', 'Why Valuable', 'Acquired Examples'],
        [
            ['Passwordless/Passkey Auth', 'Future of authentication; FIDO2 standard', 'Passage (~$24M), Passwordless.dev, Uno'],
            ['Device Trust & Zero Trust', 'Remote work security; contextual access', 'Kolide ($27M funded), Glyptodon'],
            ['Secrets Management', 'DevOps/infrastructure credentials', 'SecretHub'],
        ]
    )
    doc.add_paragraph()
    
    doc.add_heading('Tier 2: Strong Value Technologies', level=2)
    add_table(doc,
        ['Technology', 'Why Valuable', 'Acquired Examples'],
        [
            ['SaaS Access Management', 'Shadow IT discovery; app visibility', 'Trelica'],
            ['Remote Access', 'Zero-trust connections', 'Glyptodon'],
            ['Privileged Access Management', 'Enterprise security critical', 'Axiom (~$100M)'],
        ]
    )
    doc.add_paragraph()
    
    # Part 3: Contractor Vault Current State
    doc.add_heading('Part 3: Contractor Vault Current State Analysis', level=1)
    
    doc.add_heading('Current Features', level=2)
    doc.add_paragraph('• Session capture & storage with cookie encryption (Fernet/AES-128)')
    doc.add_paragraph('• Time-limited access tokens with JWT authentication')
    doc.add_paragraph('• Token claiming & injection via Chrome Extension')
    doc.add_paragraph('• Kill switch for instant revocation')
    doc.add_paragraph('• Audit trail logging and activity tracking')
    doc.add_paragraph('• Email notifications and Discord webhooks')
    doc.add_paragraph('• Contractor management system')
    
    doc.add_heading('Technical Architecture', level=2)
    add_table(doc,
        ['Component', 'Technology', 'Status'],
        [
            ['Backend API', 'FastAPI + SQLAlchemy', 'Solid'],
            ['Database', 'PostgreSQL (prod) / SQLite (dev)', 'Standard'],
            ['Extension', 'Chrome Manifest V3', 'Modern'],
            ['Encryption', 'Fernet (AES-128-CBC)', 'Secure'],
            ['Rate Limiting', 'slowapi', 'Implemented'],
            ['Audit Trail', 'Custom logging', 'Good foundation'],
        ]
    )
    doc.add_paragraph()
    
    # Part 4: Strategic Recommendations
    doc.add_heading('Part 4: Strategic Recommendations for Acquisition', level=1)
    
    doc.add_heading('Positioning Strategy', level=2)
    doc.add_paragraph(
        'Contractor Vault occupies a unique niche that aligns with current acquisition trends. '
        'Your unique angle: "Secure contractor access without password sharing"'
    )
    
    doc.add_heading('Priority 1: Passwordless/Passkey Integration', level=2)
    doc.add_paragraph('Why: Every recent acquisition (Passage, Passwordless.dev, Uno, Cotter) has this technology.')
    doc.add_paragraph('Implementation Ideas:')
    doc.add_paragraph('• Add WebAuthn support for token claiming (biometric verification)')
    doc.add_paragraph('• Integrate passkey registration for contractors')
    doc.add_paragraph('• Support for device-bound credentials')
    doc.add_paragraph('Estimated Effort: Medium (2-4 weeks)')
    doc.add_paragraph('Acquisition Value Increase: High (+$5-15M potential)')
    
    doc.add_heading('Priority 2: Device Trust & Health Monitoring', level=2)
    doc.add_paragraph('Why: Kolide was acquired specifically for this. Remote work security is critical.')
    doc.add_paragraph('Recommended Features:')
    doc.add_paragraph('• Device fingerprinting - Track which devices access tokens')
    doc.add_paragraph('• OS/Browser validation - Require specific configurations')
    doc.add_paragraph('• Location-based access - Geo-restrict token claiming')
    doc.add_paragraph('• Device health checks - Verify security posture before access')
    doc.add_paragraph('Estimated Effort: Medium-High (3-6 weeks)')
    doc.add_paragraph('Acquisition Value Increase: High (+$10-20M potential)')
    
    doc.add_heading('Priority 3: Secrets Management Extension', level=2)
    doc.add_paragraph("Why: SecretHub was acquired to add infrastructure secrets to 1Password's portfolio.")
    doc.add_paragraph('Expansion Opportunity:')
    doc.add_paragraph('• API keys management')
    doc.add_paragraph('• Database credentials')
    doc.add_paragraph('• SSH keys')
    doc.add_paragraph('• Environment variables')
    doc.add_paragraph('New Use Case: "Share temporary API access with contractors"')
    doc.add_paragraph('Estimated Effort: High (4-8 weeks)')
    doc.add_paragraph('Acquisition Value Increase: Very High (+$15-25M potential)')
    
    doc.add_heading('Priority 4: SaaS Discovery & Shadow IT', level=2)
    doc.add_paragraph('Why: Trelica was acquired (Jan 2025) specifically for this capability.')
    doc.add_paragraph('Recommended Features:')
    doc.add_paragraph('• Session Detection - Auto-detect when contractors log into new services')
    doc.add_paragraph('• App Inventory - Build a catalog of all SaaS apps contractors access')
    doc.add_paragraph('• Risk Scoring - Flag high-risk or unauthorized applications')
    doc.add_paragraph('• Access Reports - Show admins what apps are being used')
    doc.add_paragraph('Estimated Effort: Medium (3-5 weeks)')
    doc.add_paragraph('Acquisition Value Increase: High (+$10-15M potential)')
    
    # Part 5: Target Acquirer Analysis
    doc.add_heading('Part 5: Target Acquirer Analysis', level=1)
    
    add_table(doc,
        ['Company', 'Fit Score', 'Why They\'d Acquire', 'What They Lack'],
        [
            ['1Password', '★★★★★', 'Expands XAM into contractor access', 'Contractor-specific workflow'],
            ['Bitwarden', '★★★★', 'Adds enterprise contractor management', 'Time-limited session sharing'],
            ['Keeper Security', '★★★★', 'Complements connection manager', 'Contractor onboarding'],
            ['Dashlane', '★★★', 'Enterprise expansion', 'Contractor access patterns'],
            ['Okta', '★★★', 'Workforce identity extension', 'Session-based access'],
            ['CyberArk', '★★★', 'Privileged access for contractors', 'Browser-based sessions'],
        ]
    )
    doc.add_paragraph()
    
    # Part 6: Acquisition Value Estimation
    doc.add_heading('Part 6: Acquisition Value Estimation', level=1)
    
    add_table(doc,
        ['Company State', 'Estimated Range', 'Comparable'],
        [
            ['Current (session sharing only)', '$2-5M', 'Early-stage acqui-hire'],
            ['+ Passkeys', '$8-15M', 'Similar to Passage (~$24M)'],
            ['+ Device Trust', '$15-25M', 'Similar to Kolide (raised $27M)'],
            ['+ Secrets Management', '$25-40M', 'Combined value proposition'],
            ['+ SaaS Discovery', '$35-50M+', 'Full enterprise platform'],
        ]
    )
    doc.add_paragraph()
    
    # Part 7: Competitive Positioning
    doc.add_heading('Part 7: Competitive Positioning', level=1)
    
    doc.add_heading('Unique Value Proposition', level=2)
    doc.add_paragraph('"The only platform built specifically for secure contractor session access"')
    
    doc.add_heading('Key Differentiators to Emphasize', level=2)
    doc.add_paragraph('1. Session-Based Access: Share authenticated sessions, not passwords')
    doc.add_paragraph('2. Contractor-First Design: Built specifically for contractor workflows')
    doc.add_paragraph('3. Temporal Access Control: Every access has an expiration')
    doc.add_paragraph('4. Zero Password Exposure: Contractors never see the actual credentials')
    doc.add_paragraph('5. Instant Kill Switch: Revoke all access with one click')
    
    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Contractor Vault is well-positioned for acquisition with enhancements in:')
    doc.add_paragraph('1. ✓ Passwordless/Passkey Authentication - Highest value addition')
    doc.add_paragraph('2. ✓ Device Trust & Health Monitoring - Zero Trust differentiator')
    doc.add_paragraph('3. ✓ Secrets Management Extension - Platform expansion')
    doc.add_paragraph('4. ✓ SaaS Discovery & Shadow IT - Enterprise compliance value')
    doc.add_paragraph()
    doc.add_paragraph('Estimated Acquisition Range: $25M-50M with recommended enhancements')
    doc.add_paragraph('Most Likely Acquirer: 1Password (best strategic fit with XAM vision)')
    doc.add_paragraph('Timeline to Acquisition-Ready: 6-12 months with focused development')
    
    # Save the document
    doc.save('docs/ACQUISITION_STRATEGY.docx')
    print('Successfully created: docs/ACQUISITION_STRATEGY.docx')

if __name__ == '__main__':
    main()

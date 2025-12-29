"""
Convert Action Plan Markdown to DOCX
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(row_cells):
                row_cells[col_idx].text = str(cell_data)
    
    return table

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Contractor Vault: Acquisition Action Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Strategic Roadmap to Acquisition-Ready Status')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This action plan outlines the immediate, high-impact steps to position Contractor Vault '
        'for acquisition by a major password management company (1Password, Bitwarden, Keeper Security, or similar).'
    )
    p = doc.add_paragraph()
    run = p.add_run('Key Insight: ')
    run.bold = True
    p.add_run("Technology alone won't get acquired. You need paying customers + modern auth technology + proven market fit.")
    doc.add_paragraph()
    
    # Priority 1
    doc.add_heading('Priority 1: Add Passkey/WebAuthn Support', level=1)
    doc.add_paragraph('Timeline: 2-3 weeks | Impact: Critical')
    doc.add_paragraph('Every major recent acquisition (Passage, Passwordless.dev, Uno, Cotter) had passwordless technology.')
    doc.add_paragraph()
    doc.add_paragraph('Implementation:')
    doc.add_paragraph('• Add WebAuthn to token claiming flow')
    doc.add_paragraph('• Contractor enters token → Verify with Face ID/fingerprint → Session injected')
    doc.add_paragraph('• Support device-bound credentials')
    doc.add_paragraph()
    doc.add_paragraph('Transformation: "Session sharing tool" → "Passwordless contractor access platform"')
    doc.add_paragraph()
    
    # Priority 2
    doc.add_heading('Priority 2: Get 10-20 Paying Customers', level=1)
    doc.add_paragraph('Timeline: 1-3 months | Impact: Critical')
    doc.add_paragraph('Acquirers want proof of market demand. Even small revenue ($1K-5K MRR) validates the business.')
    doc.add_paragraph()
    
    doc.add_paragraph('Target Segments:')
    add_table(doc,
        ['Segment', 'Why They Need This'],
        [
            ['Digital Agencies', 'Manage contractor access to client accounts'],
            ['Startups', 'Freelancer and remote developer access'],
            ['Enterprises', 'Offshore team credential management'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph('Pricing Strategy: $29-49/month for beta access')
    doc.add_paragraph()
    doc.add_paragraph('Deliverables Needed:')
    doc.add_paragraph('☐ Landing page with pricing')
    doc.add_paragraph('☐ Stripe/payment integration')
    doc.add_paragraph('☐ Beta onboarding flow')
    doc.add_paragraph('☐ 3-5 customer case studies')
    doc.add_paragraph()
    
    # Priority 3
    doc.add_heading('Priority 3: Build Device Trust Features', level=1)
    doc.add_paragraph('Timeline: 3-4 weeks | Impact: High')
    doc.add_paragraph('Device trust is what made Kolide valuable enough for 1Password to acquire.')
    doc.add_paragraph()
    
    doc.add_paragraph('Features to Add:')
    add_table(doc,
        ['Feature', 'Description', 'Effort'],
        [
            ['Device fingerprinting', 'Track which device claimed each token', '1 week'],
            ['Device info in audit logs', 'Browser, OS, location visibility', '3 days'],
            ['Suspicious device blocking', 'Block claims from unknown devices', '1 week'],
            ['Location-based access', 'Geo-restrict token claiming', '1 week'],
        ]
    )
    doc.add_paragraph()
    
    # Path to Acquisition
    doc.add_heading('Realistic Path to Acquisition', level=1)
    add_table(doc,
        ['Timeline', 'Milestone', 'What It Proves'],
        [
            ['Now', 'Polish product, fix UX', 'Base requirement'],
            ['Month 1-2', 'Add passkeys + device tracking', 'Modern auth story'],
            ['Month 2-4', 'Get 10+ paying customers', 'Market validation'],
            ['Month 4-6', 'Case studies + security audit', 'Enterprise ready'],
            ['Month 6-12', '$10K+ MRR or 50+ customers', 'Acquisition target'],
        ]
    )
    doc.add_paragraph()
    
    # What Acquirers Look For
    doc.add_heading('What Acquirers Actually Look For', level=1)
    doc.add_paragraph('Based on analysis of Passage, Kolide, SecretHub, and Trelica acquisitions:')
    doc.add_paragraph()
    
    doc.add_heading('Must-Haves', level=2)
    doc.add_paragraph('✓ Real users paying for the product')
    doc.add_paragraph('✓ A specific niche you dominate')
    doc.add_paragraph('✓ Clean technology that integrates easily')
    doc.add_paragraph('✓ Security-first architecture (encryption, audit trails)')
    
    doc.add_heading('Nice-to-Haves', level=2)
    doc.add_paragraph('• SOC 2 compliance (or path to it)')
    doc.add_paragraph('• Enterprise SSO support')
    doc.add_paragraph('• API documentation')
    doc.add_paragraph('• Self-serve onboarding')
    doc.add_paragraph()
    
    doc.add_heading('Your Unique Niche', level=2)
    p = doc.add_paragraph()
    run = p.add_run('"Secure contractor access without sharing passwords"')
    run.italic = True
    doc.add_paragraph()
    doc.add_paragraph('This positioning sits at the intersection of:')
    doc.add_paragraph("• Session management (1Password's territory)")
    doc.add_paragraph("• Temporal access control (Kolide's value prop)")
    doc.add_paragraph('• Contractor management (enterprise compliance need)')
    doc.add_paragraph()
    
    # Action Items
    doc.add_heading('Immediate Action Items', level=1)
    
    doc.add_heading('This Week', level=2)
    doc.add_paragraph('☐ Review and prioritize passkey implementation')
    doc.add_paragraph('☐ Set up landing page with pricing')
    doc.add_paragraph('☐ Identify 20 potential beta customers to reach out to')
    
    doc.add_heading('This Month', level=2)
    doc.add_paragraph('☐ Implement basic device fingerprinting in claim endpoint')
    doc.add_paragraph('☐ Add device info to audit log display')
    doc.add_paragraph('☐ Launch beta pricing tier')
    doc.add_paragraph('☐ Start customer outreach')
    
    doc.add_heading('Next 90 Days', level=2)
    doc.add_paragraph('☐ Complete passkey/WebAuthn integration')
    doc.add_paragraph('☐ Onboard first 10 paying customers')
    doc.add_paragraph('☐ Create 2-3 customer case studies')
    doc.add_paragraph('☐ Begin SOC 2 preparation')
    doc.add_paragraph()
    
    # Success Metrics
    doc.add_heading('Success Metrics', level=1)
    add_table(doc,
        ['Metric', 'Target (90 days)', 'Target (12 months)'],
        [
            ['Paying customers', '10+', '50+'],
            ['MRR', '$500+', '$10K+'],
            ['Case studies', '2-3', '5-10'],
            ['Feature parity', 'Passkeys + Device Trust', 'Full platform'],
        ]
    )
    doc.add_paragraph()
    
    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Focus order:')
    doc.add_paragraph('1. Customers first - Prove market demand')
    doc.add_paragraph('2. Passkeys second - Modern auth is table stakes')
    doc.add_paragraph('3. Device trust third - Enterprise differentiation')
    doc.add_paragraph()
    doc.add_paragraph(
        'The companies that got acquired all had real users and a clear niche. '
        'Build that foundation, then the technology improvements become much more valuable.'
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Estimated time to acquisition-ready: 6-12 months with focused execution')
    run.italic = True
    
    # Save
    doc.save('docs/ACTION_PLAN.docx')
    print('Successfully created: docs/ACTION_PLAN.docx')

if __name__ == '__main__':
    main()
